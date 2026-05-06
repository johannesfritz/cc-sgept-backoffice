#!/usr/bin/env python3
"""
SGEPT Invoice Generator
Generates invoices by copying templates and replacing content via XML manipulation.
"""

import json
import os
import re
import shutil
import sys
import tempfile
import zipfile
from datetime import datetime, timedelta
from pathlib import Path

# Paths — script lives at claude-setup/cc-sgept-backoffice/scripts/.
# Resolve through any symlinks so the canonical locations are used regardless
# of whether the script is accessed via its real path or a symlink in jf-ceo/.
SCRIPT_PATH = Path(__file__).resolve()
REPO_DIR = SCRIPT_PATH.parent.parent          # claude-setup/cc-sgept-backoffice
JF_PRIVATE_ROOT = REPO_DIR.parent.parent      # jf-private
DATA_DIR = REPO_DIR / "config"
TEMPLATES_DIR = REPO_DIR / "templates" / "invoices"
OUTPUT_DIR = JF_PRIVATE_ROOT / "jf-ceo" / "sgept-backoffice" / "invoicing"

# Google Drive shipping destination.
# https://drive.google.com/drive/u/0/folders/19bPRghIb2L3cdxZzIattO65uM5En6dHM
# macOS: direct filesystem copy via Drive File Stream mount (canonical jf-private
#        layout uses CloudStorage/GoogleDrive-johannes.fritz@sgept.org/).
# Linux: Google Drive API via gdrive_upload.py (Phase 3).
# Override via SGEPT_GDRIVE_INVOICING env var on different machines.
GDRIVE_FOLDER_ID = "19bPRghIb2L3cdxZzIattO65uM5En6dHM"
_GDRIVE_DEFAULT = (
    f"{Path.home()}/Library/CloudStorage/GoogleDrive-johannes.fritz@sgept.org/"
    "Meine Ablage/SGEPT ORG/SGEPT admin/dbx/SGEPT/0 admin/5 invoicing"
)
GDRIVE_INVOICING_MAC = Path(os.environ.get("SGEPT_GDRIVE_INVOICING", _GDRIVE_DEFAULT))
# Back-compat alias for any external callers.
GDRIVE_INVOICING = GDRIVE_INVOICING_MAC

# Template files
TEMPLATE_STANDARD = TEMPLATES_DIR / "SGEPT - Invoice standard.docx"
TEMPLATE_NIPO_REGULAR = TEMPLATES_DIR / "SGEPT - Invoice NIPO regular.docx"
TEMPLATE_NIPO_LIBRARY = TEMPLATES_DIR / "SGEPT - Invoice NIPO academic library.docx"
TEMPLATE_NIPO_ACADEMIC = TEMPLATES_DIR / "SGEPT - Invoice NIPO academic.docx"

# Note: NIPO template values are now encoded in NIPO_LAYOUTS (see below).
# Standard invoice template values remain:
TEMPLATE_VALUES_STANDARD = {
    'company': 'Grains Australia Limited',
    'contact': 'Lachlan Evans',
    'contact_first': 'Lachlan',
    'contact_attn': 'Attn: Lachlan Evans',  # Full "Attn:" line
    'dear': 'Dear Lachlan',  # Full "Dear" line
    'street': 'C/o Ground Floor, 465 Victoria Avenue',
    'city': 'CHATSWOOD NSW 2067',
    'country': 'Australia',
    'invoice_num': '26006',
    'date_day': '15',
    'date_month': 'January',
    'date_month_short': 'Jan',
    'date_year': '2026',
    'due_day': '14',
    'due_month': 'Feb',
    'currency': 'AUD',
    'amount': '24,050.00',
    'iban': 'CH80 0078 1624 8968 1200 0',
}

def load_config():
    """Load bank account and stripe configuration."""
    with open(DATA_DIR / "bank-accounts.json") as f:
        bank_data = json.load(f)
    with open(DATA_DIR / "nipo-stripe-links.json") as f:
        stripe_data = json.load(f)
    return bank_data, stripe_data

def parse_date(date_str):
    """Parse date string and return datetime object."""
    for fmt in ["%Y-%m-%d", "%d %B %Y", "%d/%m/%Y", "%d.%m.%Y"]:
        try:
            return datetime.strptime(date_str, fmt)
        except ValueError:
            continue
    raise ValueError(f"Cannot parse date: {date_str}")

def format_amount(amount):
    """Format amount with comma thousands separator."""
    return f"{amount:,.2f}"

def get_iban_parts(iban):
    """Split IBAN into parts as they appear in templates."""
    # IBAN format: CH80 0078 1624 8968 1200 0
    parts = iban.split()
    return parts


def convert_to_pdf(docx_path):
    """Convert a .docx to .pdf next to it. Delegates to pdf_convert.convert_to_pdf
    which dispatches to docx2pdf (macOS) or libreoffice (Linux).

    Returns the Path to the resulting PDF. Raises on failure.
    """
    sys.path.insert(0, str(SCRIPT_PATH.parent))
    from pdf_convert import convert_to_pdf as _convert
    return _convert(docx_path)


def sync_to_gdrive(invoice_number, abbrev, is_nipo=False, docx_path=None):
    """Copy a local invoice (.docx + .pdf) to the Google Drive invoicing folder.

    Delegates to gdrive_upload.sync_to_gdrive which dispatches to the macOS
    Drive File Stream copy or the Linux Drive API upload.

    Returns the destination folder Path (macOS) or a remote descriptor string
    (Linux; the Drive folder URL).
    """
    sys.path.insert(0, str(SCRIPT_PATH.parent))
    from gdrive_upload import sync_to_gdrive as _sync
    return _sync(
        invoice_number=invoice_number,
        abbrev=abbrev,
        is_nipo=is_nipo,
        docx_path=docx_path,
        output_dir=OUTPUT_DIR,
    )

def replace_text_in_xml(content, old_text, new_text):
    """
    Replace text in XML, handling cases where text is split across tags.
    This is a simple replacement that works when text appears in a single <w:t> tag.
    """
    return content.replace(f">{old_text}<", f">{new_text}<")

def replace_in_docx(template_path, output_path, replacements):
    """
    Copy template and replace text in document.xml.
    Replacements is a dict of {old_text: new_text}.
    """
    with tempfile.TemporaryDirectory() as tmp_dir:
        tmp_path = Path(tmp_dir)

        # Extract docx
        with zipfile.ZipFile(template_path, 'r') as zf:
            zf.extractall(tmp_path)

        # Modify document.xml
        doc_xml = tmp_path / "word" / "document.xml"
        with open(doc_xml, 'r', encoding='utf-8') as f:
            content = f.read()

        for old, new in replacements.items():
            if old and new is not None:
                # Try various replacement patterns
                content = content.replace(f">{old}<", f">{new}<")
                content = content.replace(f">{old} <", f">{new} <")
                content = content.replace(f">{old}  <", f">{new}  <")  # Double space
                # Also try with xml:space preserve prefix
                content = content.replace(f'">{old}<', f'">{new}<')
                content = content.replace(f'">{old} <', f'">{new} <')
                content = content.replace(f'">{old}  <', f'">{new}  <')

        with open(doc_xml, 'w', encoding='utf-8') as f:
            f.write(content)

        # Repack docx
        OUTPUT_DIR.mkdir(exist_ok=True)
        with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as zf:
            for root, dirs, files in os.walk(tmp_path):
                for file in files:
                    fp = Path(root) / file
                    zf.write(fp, fp.relative_to(tmp_path))

    return output_path

# ==========================================================================
# NIPO template layouts — paragraph indices + expected template text per tier
# Uses python-docx run-level edits (robust to split runs) instead of fragile
# XML string replacement. Do NOT change template text values without first
# re-inspecting the template .docx paragraph structure.
# ==========================================================================

NIPO_LAYOUTS = {
    'regular': {
        'address': {'company': 3, 'contact': 4, 'street': 5, 'city': 6, 'country': None},
        'p_inv_num': 9,
        'p_inv_date': 10,
        'p_due_date': 11,
        'p_salutation': 13,
        'p_body': 15,
        'p_fee': 17,
        'p_total': 18,
        'p_iban': 24,
        'tpl_inv_num': '26003',
        'tpl_location_date': '13 January 2026',
        'tpl_inv_date_line': '13 Jan 2026',
        'tpl_due_date_line': '12 Feb 2026',
        'tpl_dear': 'Zachary',
        'tpl_period_start_in_body': 'January 2026',
        'tpl_period_end_in_body': 'December 2026',
        'tpl_amount': '7,000.00',
        'tpl_iban': 'CH80 0078 1624 8968 1200 0',
    },
    'academic_library': {
        'address': {'company': 3, 'contact': 4, 'street': 5, 'city': 6, 'country': None},
        'p_inv_num': 9,
        'p_inv_date': 10,
        'p_due_date': 11,
        'p_salutation': 13,
        'p_body': 15,
        'p_fee': 17,
        'p_total': 18,
        'p_iban': 24,
        'tpl_inv_num': '26005',
        'tpl_location_date': '14 January 2026',
        'tpl_inv_date_line': '14 Jan 2026',
        'tpl_due_date_line': '13 Feb 2026',
        'tpl_dear': 'Mr Lee',
        'tpl_period_start_in_body': 'January 2026',
        'tpl_period_end_in_body': 'December 2026',
        'tpl_amount': '1,250.00',
        'tpl_iban': 'CH80 0078 1624 8968 1200 0',
    },
    'academic_student': {
        'address': {'company': 2, 'contact': 3, 'street': 4, 'city': 5, 'country': 6},
        'p_inv_num': 7,
        'p_inv_date': 8,
        'p_due_date': 9,
        'p_salutation': 11,
        'p_body': 13,
        'p_fee': 15,
        'p_total': 16,
        'p_iban': 22,
        'tpl_inv_num': '25031',
        'tpl_location_date': '30 July 2025',
        'tpl_inv_date_line': '30 Jul 2025',
        'tpl_due_date_line': '29 Aug 2025',
        'tpl_dear': 'Mr. McDonald',
        'tpl_period_start_in_body': 'August 2025',
        'tpl_period_end_in_body': None,
        'tpl_amount': '500.00',
        'tpl_iban': 'CH80 0078 1624 8968 1200 0',
    },
}


def _replace_span(para, old_text, new_text):
    """
    Replace a text span (possibly split across consecutive runs) in a paragraph.
    Sets the first matching run to new_text and empties subsequent matched runs.
    Returns True if found and replaced.
    """
    if not old_text:
        return False
    runs = para.runs
    n = len(runs)
    for i in range(n):
        combined = ''
        for j in range(i, n):
            combined += runs[j].text
            if combined == old_text:
                runs[i].text = new_text
                for k in range(i + 1, j + 1):
                    runs[k].text = ''
                return True
            if not old_text.startswith(combined):
                break
    return False


def _replace_in_run_texts(para, old_text, new_text):
    """Substring replace inside each run's text (for fragments embedded in larger runs)."""
    replaced = False
    for run in para.runs:
        if old_text in run.text:
            run.text = run.text.replace(old_text, new_text)
            replaced = True
    return replaced


def _set_para_text(para, text):
    """Set paragraph to single-run text, preserving first-run formatting."""
    if not para.runs:
        para.add_run(text)
        return
    para.runs[0].text = text
    for r in para.runs[1:]:
        r.text = ''


def generate_nipo_invoice(
    invoice_number,
    invoice_date,
    currency,
    amount,
    recipient,
    tier,
    subscription_period,
    abbrev=None,
    sync=True,
):
    """
    Generate a NIPO subscription invoice using python-docx run-level edits.

    Args:
        invoice_number: 5-digit string, NO hyphen (e.g., "26015"). Hyphens are stripped.
        invoice_date: e.g., "22 January 2026" or "2026-01-22"
        currency: "CHF", "EUR", or "USD"
        amount: numeric amount (e.g., 7000)
        recipient: dict with company, name, street, city, country
        tier: "regular", "academic_library", or "academic_student"
        subscription_period: e.g., "May 2026 - April 2027"
        abbrev: short identifier used in the Drive folder name (e.g. 'IADB').
            Required when sync=True. Recommended even when sync=False so
            a re-sync via /invoice-gdrive-sync reuses the same folder name.
        sync: when True (default), generates the PDF and copies .docx+.pdf
            to the Google Drive invoicing folder with standardized names.
    """
    from docx import Document

    bank_data, _ = load_config()

    if tier == 'regular':
        template = TEMPLATE_NIPO_REGULAR
    elif tier == 'academic_library':
        template = TEMPLATE_NIPO_LIBRARY
    elif tier == 'academic_student':
        template = TEMPLATE_NIPO_ACADEMIC
    else:
        raise ValueError(f"Unknown tier: {tier}")

    if not template.exists():
        raise FileNotFoundError(f"Template not found: {template}")

    layout = NIPO_LAYOUTS[tier]

    inv_num = str(invoice_number).replace('-', '').strip()
    dt = parse_date(invoice_date)
    due_dt = dt + timedelta(days=30)

    period_parts = [p.strip() for p in subscription_period.split('-')]
    period_start = period_parts[0] if period_parts else subscription_period
    period_end = period_parts[1] if len(period_parts) > 1 else ''

    contact_name = recipient.get('name', '') or ''
    contact_first = contact_name.split()[0] if contact_name else ''
    iban = bank_data['accounts'][currency]['iban']

    location_date_str = f"{dt.day} {dt.strftime('%B')} {dt.year}"
    inv_date_line_str = f"{dt.day} {dt.strftime('%b')} {dt.year}"
    due_date_line_str = f"{due_dt.day} {due_dt.strftime('%b')} {due_dt.year}"
    amount_str = format_amount(amount)

    doc = Document(template)
    paras = list(doc.paragraphs)

    # Address block
    addr = layout['address']
    _set_para_text(paras[addr['company']], recipient.get('company', ''))
    _set_para_text(paras[addr['contact']], contact_name)
    _set_para_text(paras[addr['street']], recipient.get('street', ''))
    _set_para_text(paras[addr['city']], recipient.get('city', ''))
    if addr['country'] is not None:
        _set_para_text(paras[addr['country']], recipient.get('country', ''))

    # Invoice number paragraph (contains invoice number AND right-side location date)
    p_inv = paras[layout['p_inv_num']]
    if not _replace_span(p_inv, layout['tpl_inv_num'], inv_num):
        raise RuntimeError(
            f"Could not replace template invoice number '{layout['tpl_inv_num']}' in P[{layout['p_inv_num']}] "
            f"of {template.name}. Template run structure may have changed; re-inspect the .docx."
        )
    if not _replace_span(p_inv, layout['tpl_location_date'], location_date_str):
        raise RuntimeError(
            f"Could not replace template location date '{layout['tpl_location_date']}' in P[{layout['p_inv_num']}] "
            f"of {template.name}."
        )

    # Invoice date line
    p_date = paras[layout['p_inv_date']]
    if not _replace_span(p_date, layout['tpl_inv_date_line'], inv_date_line_str):
        raise RuntimeError(
            f"Could not replace template invoice date '{layout['tpl_inv_date_line']}' in P[{layout['p_inv_date']}] "
            f"of {template.name}."
        )

    # Due date line
    p_due = paras[layout['p_due_date']]
    if not _replace_span(p_due, layout['tpl_due_date_line'], due_date_line_str):
        raise RuntimeError(
            f"Could not replace template due date '{layout['tpl_due_date_line']}' in P[{layout['p_due_date']}] "
            f"of {template.name}."
        )

    # Salutation (regular → first name; academic templates → full form 'Mr X')
    p_dear = paras[layout['p_salutation']]
    dear_replacement = contact_first if tier == 'regular' else contact_name
    _replace_span(p_dear, layout['tpl_dear'], dear_replacement)

    # Body paragraph: period start/end
    p_body = paras[layout['p_body']]
    if layout['tpl_period_start_in_body']:
        _replace_span(p_body, layout['tpl_period_start_in_body'], period_start)
    if layout['tpl_period_end_in_body']:
        _replace_span(p_body, layout['tpl_period_end_in_body'], period_end)

    # Fee + Total lines: amount + currency
    for p_idx in [layout['p_fee'], layout['p_total']]:
        p = paras[p_idx]
        if not _replace_span(p, layout['tpl_amount'], amount_str):
            _replace_in_run_texts(p, layout['tpl_amount'], amount_str)
        _replace_in_run_texts(p, 'CHF', currency)

    # IBAN
    p_iban = paras[layout['p_iban']]
    if not _replace_span(p_iban, layout['tpl_iban'], iban):
        _replace_in_run_texts(p_iban, layout['tpl_iban'], iban)

    # Output path
    company_short = recipient.get('company', contact_name or 'Client')
    if len(company_short) > 25:
        company_short = company_short[:25].rsplit(' ', 1)[0]
    folder_name = f"{dt.strftime('%y%m%d')} NIPO {company_short} {inv_num}"
    output_subdir = OUTPUT_DIR / folder_name
    output_subdir.mkdir(exist_ok=True)

    company_slug = re.sub(r'[^\w\s-]', '', recipient.get('company', contact_name or 'Client'))
    company_slug = company_slug.replace(' ', '-')[:30]
    filename = f"Invoice-{inv_num}-{company_slug}-NIPO.docx"
    output_path = output_subdir / filename

    doc.save(output_path)

    print(f"\n{'='*60}")
    print(f"NIPO Invoice Generated: {output_path.name}")
    print(f"{'='*60}")
    print(f"Template: {template.name}")
    print(f"Recipient: {recipient.get('company')}")
    print(f"Contact: {contact_name}")
    print(f"Invoice #: {inv_num}")
    print(f"Date: {dt.strftime('%d %B %Y')}")
    print(f"Due: {due_dt.strftime('%d %B %Y')}")
    print(f"Amount: {currency} {amount_str}")
    print(f"IBAN: {iban}")
    print(f"Period: {period_start} - {period_end}" if period_end else f"Period: {period_start}")
    print(f"\nOutput: {output_path}")

    if sync:
        if not abbrev:
            raise ValueError("sync=True requires abbrev (e.g. 'IADB', 'ERCO').")
        sync_to_gdrive(inv_num, abbrev, is_nipo=True, docx_path=output_path)

    return output_path

def generate_standard_invoice(
    invoice_number,
    invoice_date,
    currency,
    recipient,
    subject_line,
    intro_paragraph,
    items,
    items_type,
    total,
    vat_note,
    abbrev=None,
    sync=True,
):
    """
    Generate a standard professional services invoice using python-docx.

    Args:
        invoice_number: e.g., "26008"
        invoice_date: e.g., "12 February 2026"
        currency: "CHF", "EUR", or "USD"
        recipient: dict with company, name, street, city, country
        subject_line: project title, e.g., "GTA Policy Intelligence Dashboard"
        intro_paragraph: full intro text starting with agreement reference
        items: list of dicts:
            deliverable: [{"description": "...", "amount": 12000}]
            hours: [{"description": "J. Fritz", "hours": 45, "rate": 300, "total": 13500}]
        items_type: "deliverable" or "hours"
        total: total invoice amount (number)
        vat_note: e.g., "No VAT is applied."
        abbrev: short identifier used in the Drive folder name (e.g. 'ERCO').
            Required when sync=True.
        sync: when True (default), generates the PDF and copies .docx+.pdf
            to the Google Drive invoicing folder with standardized names.
    """
    from docx import Document
    from docx.shared import Emu
    from copy import deepcopy
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    bank_data, _ = load_config()

    if not TEMPLATE_STANDARD.exists():
        raise FileNotFoundError(f"Template not found: {TEMPLATE_STANDARD}")

    dt = parse_date(invoice_date)
    due_dt = dt + timedelta(days=30)
    inv_num = invoice_number.replace("-", "")
    iban = bank_data['accounts'][currency]['iban']
    contact_name = recipient.get('name', '')
    contact_first = contact_name.split()[0] if contact_name else ''

    doc = Document(TEMPLATE_STANDARD)
    paras = list(doc.paragraphs)

    SMALL_FONT = Emu(101600)  # ~8pt for metadata labels

    # --- Helpers ---
    def clear_runs(para):
        for r in list(para._element.findall(qn('w:r'))):
            para._element.remove(r)

    def set_text(para, text, bold=None):
        clear_runs(para)
        run = para.add_run(text)
        if bold is not None:
            run.bold = bold

    def delete_para(para):
        para._element.getparent().remove(para._element)

    def make_para_el(text, ref_el):
        """Create paragraph XML element with text, copying format from ref."""
        new_p = deepcopy(ref_el)
        for r in new_p.findall(qn('w:r')):
            new_p.remove(r)
        r_el = OxmlElement('w:r')
        t_el = OxmlElement('w:t')
        t_el.text = text
        t_el.set(qn('xml:space'), 'preserve')
        r_el.append(t_el)
        new_p.append(r_el)
        return new_p

    def make_bold_para_el(label, value, ref_el):
        """Create paragraph with bold label + bold value (for total line)."""
        new_p = deepcopy(ref_el)
        for r in new_p.findall(qn('w:r')):
            new_p.remove(r)
        for text in [label, value]:
            r_el = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            b_el = OxmlElement('w:b')
            rPr.append(b_el)
            r_el.append(rPr)
            t_el = OxmlElement('w:t')
            t_el.text = text
            t_el.set(qn('xml:space'), 'preserve')
            r_el.append(t_el)
            new_p.append(r_el)
        return new_p

    # --- Save references to all template paragraphs ---
    # Template structure (by index):
    #   P[3-7]: Address | P[9-11]: Metadata | P[13]: Dear
    #   P[15]: Intro | P[17]: "Service Details:" | P[18]: Col headers
    #   P[19-21]: Item rows | P[22]: Total
    #   P[24]: "Note:" | P[25-27]: Note lines
    #   P[29]: Bank intro | P[31-34]: Bank details | P[36]: Closing
    p_addr = [paras[3], paras[4], paras[5], paras[6], paras[7]]
    p_inv_num = paras[9]
    p_inv_date = paras[10]
    p_due_date = paras[11]
    p_salutation = paras[13]
    p_intro = paras[15]
    p_svc_header = paras[17]
    p_col_headers = paras[18]
    p_item_rows = [paras[19], paras[20], paras[21]]
    p_total = paras[22]
    p_note_header = paras[24]
    p_note_lines = [paras[25], paras[26], paras[27]]
    p_iban = paras[33]

    # Save item row XML for cloning before we delete them
    item_row_el = deepcopy(p_item_rows[0]._element)

    # === Address block ===
    set_text(p_addr[0], recipient.get('company', ''))
    set_text(p_addr[1], f"Attn: {contact_name}")
    set_text(p_addr[2], recipient.get('street', ''))
    set_text(p_addr[3], recipient.get('city', ''))
    set_text(p_addr[4], recipient.get('country', ''))

    # === Invoice number + location date ===
    clear_runs(p_inv_num)
    r = p_inv_num.add_run(f"Invoice number:  {inv_num}\t")
    r.font.size = SMALL_FONT
    p_inv_num.add_run(f"St.Gallen, {dt.day} {dt.strftime('%B')} {dt.year}")

    # === Invoice date + subject ===
    clear_runs(p_inv_date)
    r = p_inv_date.add_run(f"Invoice date:         {dt.day} {dt.strftime('%b')} {dt.year}\t")
    r.font.size = SMALL_FONT
    r2 = p_inv_date.add_run(f"Invoice: {subject_line}")
    r2.font.name = 'Roboto Medium'

    # === Due date ===
    clear_runs(p_due_date)
    r = p_due_date.add_run(
        f"Due date:               {due_dt.day} {due_dt.strftime('%b')} {due_dt.year}"
    )
    r.font.size = SMALL_FONT

    # === Salutation ===
    set_text(p_salutation, f"Dear {contact_first}")

    # === Intro paragraph ===
    set_text(p_intro, intro_paragraph)

    # === Column headers ===
    if items_type == 'hours':
        set_text(p_col_headers,
                 f"Description\t\t\tHours\tRate ({currency})\t\tTotal ({currency})")
    else:
        set_text(p_col_headers,
                 f"Description\t\t\t\t\t\tAmount ({currency})")

    # === Item rows: delete old, insert new ===
    for p in p_item_rows:
        delete_para(p)

    for item in items:
        if items_type == 'hours':
            text = (f"{item['description']}\t\t\t"
                    f"{item.get('hours', '')}\t"
                    f"{format_amount(item.get('rate', 0))}\t\t"
                    f"{format_amount(item.get('total', 0))}")
        else:
            text = f"{item['description']}\t\t\t\t\t\t{format_amount(item['amount'])}"
        new_p = make_para_el(text, item_row_el)
        p_total._element.addprevious(new_p)

    # === Total line ===
    clear_runs(p_total)
    r = p_total.add_run("Invoice total:\t\t\t\t\t\t")
    r.bold = True
    r = p_total.add_run(format_amount(total))
    r.bold = True

    # === Notes: replace with VAT note ===
    for p in p_note_lines:
        delete_para(p)
    set_text(p_note_header, vat_note)

    # === IBAN (update for currency) ===
    clear_runs(p_iban)
    p_iban.add_run("IBAN:")
    p_iban.add_run(f"\t\t{iban}")

    # === Save output — into subfolder: YYMMDD Subject Recipient NNNNN/ ===
    company_short = recipient.get('company', 'Client')
    if len(company_short) > 25:
        company_short = company_short[:25].rsplit(' ', 1)[0]
    # Use first word of subject_line as descriptor (e.g. "GTA-DB", "Dashboard")
    descriptor = subject_line.split()[0] if subject_line else 'Invoice'
    folder_name = f"{dt.strftime('%y%m%d')} {descriptor} {company_short} {inv_num}"
    output_subdir = OUTPUT_DIR / folder_name
    output_subdir.mkdir(exist_ok=True)

    company_slug = re.sub(r'[^\w\s-]', '', recipient.get('company', 'Client'))
    company_slug = company_slug.replace(' ', '-')[:30]
    filename = f"Invoice-{invoice_number}-{company_slug}.docx"
    output_path = output_subdir / filename
    doc.save(output_path)

    print(f"\n{'='*60}")
    print(f"Standard Invoice Generated: {output_path.name}")
    print(f"{'='*60}")
    print(f"Recipient: {recipient.get('company')}")
    print(f"Contact: {recipient.get('name')}")
    print(f"Invoice #: {invoice_number}")
    print(f"Date: {dt.strftime('%d %B %Y')}")
    print(f"Due: {due_dt.strftime('%d %B %Y')}")
    print(f"Amount: {currency} {format_amount(total)}")
    print(f"IBAN: {iban}")
    print(f"Subject: {subject_line}")
    print(f"\nOutput: {output_path}")

    if sync:
        if not abbrev:
            raise ValueError("sync=True requires abbrev (e.g. 'ERCO', 'IADB').")
        sync_to_gdrive(inv_num, abbrev, is_nipo=False, docx_path=output_path)

    return output_path


def generate_from_spec(spec):
    """Drive generation from a spec dict (same schema as CLI --spec-file).

    Required keys: type ('nipo'|'standard'), invoice_number, invoice_date,
    currency, recipient, abbrev. Additional per-type fields described in
    knowledge/invoice-spec-schema.md.
    """
    inv_type = spec.get('type', '').lower()
    common = dict(
        invoice_number=spec['invoice_number'],
        invoice_date=spec['invoice_date'],
        currency=spec['currency'],
        recipient=spec['recipient'],
        abbrev=spec['abbrev'],
        sync=spec.get('sync', True),
    )
    if inv_type == 'nipo':
        return generate_nipo_invoice(
            tier=spec['tier'],
            amount=spec['amount'],
            subscription_period=spec['subscription_period'],
            **common,
        )
    if inv_type == 'standard':
        return generate_standard_invoice(
            subject_line=spec['subject_line'],
            intro_paragraph=spec['intro_paragraph'],
            items=spec['items'],
            items_type=spec['items_type'],
            total=spec['total'],
            vat_note=spec.get('vat_note', 'No VAT is applied.'),
            **common,
        )
    raise ValueError(f"Unknown invoice type: {inv_type!r}. Expected 'nipo' or 'standard'.")


def _cli():
    import argparse
    ap = argparse.ArgumentParser(description="SGEPT Invoice Generator")
    ap.add_argument('--spec-file', type=Path,
                    help="Path to JSON spec file (see knowledge/invoice-spec-schema.md).")
    ap.add_argument('--no-sync', action='store_true',
                    help="Skip Drive upload (overrides spec.sync).")
    args = ap.parse_args()

    if args.spec_file:
        with open(args.spec_file) as f:
            spec = json.load(f)
        if args.no_sync:
            spec['sync'] = False
        out = generate_from_spec(spec)
        print(f"INVOICE_OUTPUT: {out}")
        return

    # No args: show status
    print("SGEPT Invoice Generator")
    print(f"Templates: {TEMPLATES_DIR}")
    print(f"Output: {OUTPUT_DIR}")
    print("\nAvailable templates:")
    for t in [TEMPLATE_STANDARD, TEMPLATE_NIPO_REGULAR, TEMPLATE_NIPO_LIBRARY, TEMPLATE_NIPO_ACADEMIC]:
        status = "OK" if t.exists() else "MISSING"
        print(f"  [{status}] {t.name}")


if __name__ == "__main__":
    _cli()
